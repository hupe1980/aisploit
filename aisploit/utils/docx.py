from dataclasses import dataclass
from typing import IO, Union

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt


@dataclass
class Docx:
    docx: Union[str, IO[bytes], None] = None

    def __post_init__(self):
        if self.docx:
            if isinstance(self.docx, str):
                self.document = Document(self.docx)
            else:
                self.document = Document(self.docx)
        else:
            self.document = Document()

    def add_paragraph(
        self,
        *,
        text: str,
        style: str | None = None,
        alignment: str | None = None,
        hidden: bool = False,
        font_size: float | None = None
    ) -> Document:
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        if style:
            run.style = style
        if alignment:
            paragraph.alignment = alignment
        if font_size:
            run.font.size = Pt(font_size)
        if hidden:
            run.font.hidden = True  # Set the font to be hidden
            self._set_hidden_property(paragraph)
        return self.document

    def add_heading(
        self,
        *,
        text: str,
        level: int = 1,
        alignment: str | None = None,
        hidden: bool = False,
        font_size: float | None = None
    ) -> Document:
        heading = self.document.add_heading(level)
        run = heading.add_run(text)
        if alignment:
            heading.alignment = alignment
        if font_size:
            run.font.size = Pt(font_size)
        if hidden:
            run.font.hidden = True  # Set the font to be hidden
            self._set_hidden_property(heading)
        return self.document

    def save(self, filename: str) -> None:
        self.document.save(filename)

    def _set_hidden_property(self, element):
        pPr = OxmlElement('w:pPr')  # paragraph property
        rPr = OxmlElement('w:rPr')  # run property
        v = OxmlElement('w:vanish')  # hidden
        rPr.append(v)
        pPr.append(rPr)
        element._p.append(pPr)


class DocxTemplate(Docx):
    pass
